import os
import re
import uuid
from io import BytesIO
from copy import deepcopy
from xml.sax.saxutils import escape as xml_escape
from typing import Dict, List, Tuple, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table
from docx.table import _Row as Row, _Cell as Cell
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REGION_TYPE_KEYWORDS = {
    '基本信息': ['基本信息', '学科', '年级', '课题', '课时', '授课人', '班级', '教材版本', '课型'],
    '教学目标': ['教学目标', '三维目标', '核心素养', '知识与技能', '过程与方法', '情感态度'],
    '教学重难点': ['教学重难点', '重点', '难点', '教学重点', '教学难点', '重难点'],
    '学情分析': ['学情分析', '学生分析', '学习情况', '学习者分析'],
    '教学方法': ['教学方法', '教法', '学法', '教学策略', '教学模式'],
    '教学过程': ['教学过程', '教学环节', '教学步骤', '教学活动', '课堂流程'],
    '导入': ['导入', '情境导入', '复习导入', '激趣导入'],
    '新授': ['新授', '新知讲授', '讲授新知', '知识讲解'],
    '巩固练习': ['巩固练习', '课堂练习', '练习', '训练'],
    '小结': ['小结', '课堂小结', '总结', '归纳'],
    '作业布置': ['作业布置', '课后作业', '作业', '练习作业'],
    '板书设计': ['板书设计', '板书', '板演'],
    '教学反思': ['教学反思', '反思', '教学后记', '教后反思'],
    '时间分配': ['时间分配', '课时安排', '时间安排', '各环节时间'],
    '教具准备': ['教具准备', '教学准备', '教具', '学具'],
    '教学媒体': ['教学媒体', '多媒体', '课件', '信息技术'],
}

PLACEHOLDER_PATTERNS = [
    re.compile(r'_{3,}', re.UNICODE),
    re.compile(r'\[\s*___+\s*\]', re.UNICODE),
    re.compile(r'\[\s*待填写\s*\]', re.UNICODE),
    re.compile(r'\[\s*填写说明\s*\]', re.UNICODE),
    re.compile(r'\[\s*内容\s*\]', re.UNICODE),
    re.compile(r'\[\s*\]\s*'),
    re.compile(r'【\s*填写.*?】', re.UNICODE),
    re.compile(r'（\s*填写.*?）', re.UNICODE),
]


class DocxStructureParser:
    def __init__(self, docx_bytes: bytes):
        self.docx_bytes = docx_bytes
        self.structure_map: Dict[str, Dict] = {}
        self.pending_confirmations: List[Dict] = []
        self.region_counter = 0

    def parse(self) -> Tuple[Dict, List[Dict]]:
        # 修复 M5：用 try/finally 显式关闭 BytesIO，确保资源释放
        # （解析出的 lxml 元素已载入内存，关闭 bio 不影响 structure_map 中的 _element_ref）
        bio = BytesIO(self.docx_bytes)
        try:
            self.doc = Document(bio)
            self._parse_paragraphs()
            self._parse_tables()
            self._parse_textboxes()
            self._apply_smart_inference()
            self._build_pending_confirmations()
        finally:
            bio.close()
        return self.structure_map, self.pending_confirmations

    def _parse_paragraphs(self):
        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            region_type = self._detect_region_type(text)
            is_placeholder = self._is_placeholder(text)

            region_info = {
                'type': 'paragraph',
                'index': idx,
                'location': f'段落 {idx + 1}',
                'text': text[:200],
                'full_text': text,
                'region_type': region_type,
                'is_placeholder': is_placeholder,
                'style': self._extract_paragraph_style(para),
                'needs_confirmation': False,
                '_raw_type': region_type,
            }

            region_id = f'region_{self.region_counter}'
            self.region_counter += 1
            self.structure_map[region_id] = region_info

    def _parse_tables(self):
        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if not cell_text:
                        continue

                    region_type = self._detect_region_type(cell_text)
                    is_placeholder = self._is_placeholder(cell_text)

                    cell_info = self._extract_cell_info(table, row_idx, col_idx, cell)

                    region_info = {
                    'type': 'table_cell',
                    'table_index': table_idx,
                    'row_index': row_idx,
                    'col_index': col_idx,
                    'location': f'表格 {table_idx + 1}, 单元格 ({row_idx + 1}, {col_idx + 1})',
                    'text': cell_text[:200],
                    'full_text': cell_text,
                    'region_type': region_type,
                    'is_placeholder': is_placeholder,
                    'cell_info': cell_info,
                    'style': self._extract_cell_style(cell),
                    'needs_confirmation': False,
                    '_raw_type': region_type,
                }

                    region_id = f'region_{self.region_counter}'
                    self.region_counter += 1
                    self.structure_map[region_id] = region_info

    def _parse_textboxes(self):
        """解析文档中的文本框（VML w:pict/v:textbox 和 DrawingML wps:txbx）"""
        # 修复 M4：精确匹配 textbox/txbxContent 本地标签名，命中后跳过其子树，
        # 避免 v:textbox 与其内部 w:txbxContent 被重复采集
        TEXTBOX_TAGS = {'textbox', 'txbxcontent'}
        # 1. VML 文本框：<w:pict> 内的 <v:textbox>
        collected_elements = []  # 记录已采集的文本框元素，用于跳过其子树
        for pict in self.doc.element.body.iter(qn('w:pict')):
            for txbx in pict.iter():
                # 取本地标签名（去除命名空间）并小写，精确匹配
                local_tag = txbx.tag.split('}')[-1].lower() if '}' in txbx.tag else txbx.tag.lower()
                if local_tag not in TEXTBOX_TAGS:
                    continue
                # 跳过已采集元素的子树，避免重复采集同一文本框
                if any(txbx in collected.iter() for collected in collected_elements):
                    continue
                collected_elements.append(txbx)
                texts = []
                for t in txbx.iter(qn('w:t')):
                    if t.text:
                        texts.append(t.text)
                full_text = ''.join(texts).strip()
                if not full_text:
                    continue
                region_type = self._detect_region_type(full_text)
                is_placeholder = self._is_placeholder(full_text)
                region_id = f'region_{self.region_counter}'
                self.region_counter += 1
                self.structure_map[region_id] = {
                    'type': 'textbox',
                    'location': '文本框',
                    'full_text': full_text,
                    'region_type': region_type,
                    'is_placeholder': is_placeholder,
                    '_raw_type': region_type,
                    '_element_ref': txbx,
                }

        # 2. DrawingML 文本框：<wps:txbx> 内的 <w:txbxContent>
        for drawing in self.doc.element.body.iter(qn('w:drawing')):
            for txbx in drawing.iter():
                tag = txbx.tag
                if 'txbx' in tag.lower() and 'txbxContent' not in tag:
                    txbxContent = None
                    for child in txbx.iter():
                        if 'txbxContent' in child.tag:
                            txbxContent = child
                            break
                    if txbxContent is None:
                        continue
                    texts = []
                    for t in txbxContent.iter(qn('w:t')):
                        if t.text:
                            texts.append(t.text)
                    full_text = ''.join(texts).strip()
                    if not full_text:
                        continue
                    region_type = self._detect_region_type(full_text)
                    is_placeholder = self._is_placeholder(full_text)
                    region_id = f'region_{self.region_counter}'
                    self.region_counter += 1
                    self.structure_map[region_id] = {
                        'type': 'textbox',
                        'location': '文本框',
                        'full_text': full_text,
                        'region_type': region_type,
                        'is_placeholder': is_placeholder,
                        '_raw_type': region_type,
                        '_element_ref': txbxContent,
                    }

    def _detect_region_type(self, text: str) -> str:
        for region_type, keywords in REGION_TYPE_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text:
                    return region_type
        # 内容特征推断（无关键词命中时）
        return self._infer_by_content_features(text)

    def _infer_by_content_features(self, text: str) -> str:
        t = text.lower()
        scores = {}

        def add_score(rtype, pts):
            scores[rtype] = scores.get(rtype, 0) + pts

        # 时间/课时（强特征）
        if re.search(r'\d+\s*[分钟分节课时]|第\d+[节课时]', text):
            add_score('时间分配', 5)
            add_score('教学过程', 1)
        # 基本信息（严格匹配：带冒号或完整词）
        if re.search(r'^(学科|年级|课题|课时|课型|授课人|班级|科目|教材)\s*[：:]', text) or text in ('学科', '年级', '课题', '课时', '课型', '授课人', '班级', '科目', '教材'):
            add_score('基本信息', 5)
        # 学情相关
        for k in ['学情分析', '学生分析', '学习情况', '学习者分析', '学情']:
            if k in t:
                add_score('学情分析', 4 if '学情' in k else 2)
        # 方法相关
        for k in ['讲授法', '讨论法', '演示法', '探究法', '小组合作', '启发式', '任务驱动', '教学方法', '教学策略']:
            if k in t:
                add_score('教学方法', 3)
        # 教具/媒体
        for k in ['多媒体', 'ppt', '课件', '投影仪', '实验器材', '教具', '学具', '教学媒体', '教具准备']:
            if k in t:
                add_score('教学媒体', 2)
                add_score('教具准备', 1)
        # 目标相关
        if '教学目标' in t or '学习目标' in t:
            add_score('教学目标', 5)
        elif any(k in t for k in ['目标', '目的', '要求']):
            add_score('教学目标', 3 if len(text) < 30 else 1)
        # 重难点相关
        if '教学重难点' in t or '教学重点' in t or '教学难点' in t:
            add_score('教学重难点', 5)
        elif any(k in t for k in ['重点', '难点']):
            add_score('教学重难点', 3 if len(text) < 30 else 1)
        # 作业相关
        for k in ['作业布置', '课后作业', '作业', '习题']:
            if k in t:
                add_score('作业布置', 3 if '作业' in k else 2)
        # 板书
        if '板书设计' in t or '板书' in t:
            add_score('板书设计', 4)
        # 反思
        if '教学反思' in t or '教后反思' in t or '反思' in t:
            add_score('教学反思', 4)
        # 导入
        if any(k in t for k in ['导入', '引入', '激趣']):
            add_score('导入', 3)
        # 小结
        if any(k in t for k in ['课堂小结', '小结', '总结', '归纳']):
            add_score('小结', 3 if '小结' in t or '总结' in t else 2)
        # 练习/巩固
        for k in ['巩固练习', '课堂练习', '练习', '巩固', '训练', '例题']:
            if k in t:
                add_score('巩固练习', 3 if '练习' in k else 2)
        # 新授/讲授
        for k in ['新授', '新知讲授', '讲授新知']:
            if k in t:
                add_score('新授', 4)
        if '讲授' in t or '讲解' in t or '新知' in t:
            add_score('新授', 1)
        # 教学过程通用词
        for k in ['教学过程', '教学环节', '教学步骤', '教学活动', '课堂流程']:
            if k in t:
                add_score('教学过程', 4)
        if '教学' in t or '环节' in t or '活动' in t or '过程' in t:
            add_score('教学过程', 1)

        # 长文本：教学过程权重更高，子环节类权重降低
        if len(text) > 45:
            add_score('教学过程', 6)
            # 统计命中了多少种子环节类型
            sub_types = ['新授', '导入', '小结', '巩固练习', '作业布置', '板书设计', '教学反思']
            hit_sub = sum(1 for s in sub_types if scores.get(s, 0) > 0)
            # 命中多种子环节 → 更可能是综合性的教学过程段落
            if hit_sub >= 2:
                add_score('教学过程', hit_sub * 2)
            # 长文本中，子环节类型只作为辅助信号，大幅降权
            for rtype in sub_types:
                if rtype in scores and scores[rtype] > 1:
                    scores[rtype] = 1

        if not scores:
            return '待确认'

        best_type = max(scores, key=scores.get)
        best_score = scores[best_type]

        # 得分太低且文本短，标记为待确认
        if best_score < 2 and len(text) < 20:
            return '待确认'

        return best_type

    def _apply_smart_inference(self):
        """第二轮：基于位置和上下文做推断，并设置默认回退。"""
        items = list(self.structure_map.items())
        # 1. 段落位置推断
        paragraph_items = [(rid, r) for rid, r in items if r['type'] == 'paragraph']
        for i, (rid, region) in enumerate(paragraph_items):
            if region['_raw_type'] != '待确认':
                continue
            text = region['full_text']
            # 第一个段落通常是标题或基本信息
            if i == 0 and len(text) < 30:
                region['region_type'] = '基本信息'
                continue
            # 前3个非标题段落，如果是短文本，可能是教学目标/学情
            if i < 3 and len(text) < 100:
                if '目标' in text or '素养' in text:
                    region['region_type'] = '教学目标'
                elif '学生' in text or '班级' in text:
                    region['region_type'] = '学情分析'
                elif '方法' in text or '策略' in text:
                    region['region_type'] = '教学方法'
                else:
                    region['region_type'] = '教学过程'
                continue
            # 默认回退
            region['region_type'] = '教学过程'

        # 2. 表格位置推断
        table_items = [(rid, r) for rid, r in items if r['type'] == 'table_cell']
        # 按表格分组
        tables_dict = {}
        for rid, region in table_items:
            tidx = region['table_index']
            tables_dict.setdefault(tidx, []).append((rid, region))
        for tidx, cells in tables_dict.items():
            # 按行、列排序
            cells.sort(key=lambda x: (x[1]['row_index'], x[1]['col_index']))
            # 第一行通常是表头
            first_row_idx = min(c[1]['row_index'] for c in cells)
            for rid, region in cells:
                if region['_raw_type'] != '待确认':
                    continue
                row_idx = region['row_index']
                col_idx = region['col_index']
                text = region['full_text']
                # 表头行推断
                if row_idx == first_row_idx:
                    if '时间' in text or '环节' in text or '步骤' in text:
                        region['region_type'] = '教学过程'
                    elif '目标' in text:
                        region['region_type'] = '教学目标'
                    elif '重难' in text:
                        region['region_type'] = '教学重难点'
                    else:
                        region['region_type'] = '教学过程'
                    continue
                # 第一列通常是标签列
                first_col_idx = min(c[1]['col_index'] for c in cells if c[1]['row_index'] == row_idx)
                if col_idx == first_col_idx:
                    if '导入' in text:
                        region['region_type'] = '导入'
                    elif '新授' in text or '讲授' in text or '讲解' in text:
                        region['region_type'] = '新授'
                    elif '练习' in text or '巩固' in text:
                        region['region_type'] = '巩固练习'
                    elif '小结' in text or '总结' in text:
                        region['region_type'] = '小结'
                    elif '作业' in text:
                        region['region_type'] = '作业布置'
                    elif '板书' in text:
                        region['region_type'] = '板书设计'
                    elif '目标' in text:
                        region['region_type'] = '教学目标'
                    elif '重难' in text:
                        region['region_type'] = '教学重难点'
                    elif '反思' in text:
                        region['region_type'] = '教学反思'
                    else:
                        region['region_type'] = '教学过程'
                    continue
                # 非标签列的单元格，默认跟随同列或同行的已知类型
                region['region_type'] = '教学过程'

    def _build_pending_confirmations(self):
        """仅保留极少数极度模糊、AI 也无法可靠推断的区域。"""
        self.pending_confirmations = []
        for region_id, region in self.structure_map.items():
            # 只有同时满足以下条件的才需要人工确认：
            # 1. 原始检测为待确认；2. 不是占位符；3. 文本极短（<4字）或完全无意义；4. 推断后仍为待确认
            if region.get('_raw_type') == '待确认' and not region.get('is_placeholder', False) and region.get('region_type') == '待确认':
                text = region.get('full_text', '').strip()
                if len(text) < 4 or text in ['无', '暂无', '/']:
                    self.pending_confirmations.append({
                        'region_id': region_id,
                        'location': region.get('location', ''),
                        'content_preview': text[:100],
                        'detected_type': region['region_type'],
                        'suggestions': self._get_suggestions(text),
                    })
                    region['needs_confirmation'] = True

    def _is_placeholder(self, text: str) -> bool:
        for pattern in PLACEHOLDER_PATTERNS:
            if pattern.search(text):
                return True
        return False

    def _extract_paragraph_style(self, para) -> Dict:
        color = None
        if para.style.font and para.style.font.color:
            try:
                color = para.style.font.color.rgb
            except (ValueError, TypeError):
                color = None
        style = {
            'alignment': str(para.alignment),
            'font_name': para.style.font.name if para.style.font else None,
            'font_size': para.style.font.size.pt if para.style.font and para.style.font.size else None,
            'bold': para.style.font.bold if para.style.font else None,
            'italic': para.style.font.italic if para.style.font else None,
            'color': color,
            'line_spacing': para.paragraph_format.line_spacing,
            'space_before': para.paragraph_format.space_before,
            'space_after': para.paragraph_format.space_after,
        }
        return style

    def _extract_cell_style(self, cell) -> Dict:
        first_para = cell.paragraphs[0] if cell.paragraphs else None
        style = {}
        if first_para:
            style.update(self._extract_paragraph_style(first_para))

        tc = cell._tc
        tcPr = tc.tcPr
        if tcPr is not None:
            style['width'] = tcPr.tcW.val if tcPr.tcW else None
            style['vertical_align'] = tcPr.vAlign.val if tcPr.vAlign else None

        return style

    def _extract_cell_info(self, table: Table, row_idx: int, col_idx: int, cell: Cell) -> Dict:
        tc = cell._tc
        tcPr = tc.tcPr

        rowspan = 1
        vMerge = tcPr.find(qn('w:vMerge')) if tcPr is not None else None
        if vMerge is not None:
            # 注意 L2：简化处理，仅支持 rowspan=2（合并起点 w:val="restart" 与续点无 val）
            # TODO: 完整实现应累计 continue 数量计算真实 rowspan
            if vMerge.get(qn('w:val')) is None:
                rowspan = 1
            else:
                rowspan = 2

        colspan = 1
        gridSpan = tcPr.find(qn('w:gridSpan')) if tcPr is not None else None
        if gridSpan is not None:
            colspan = int(gridSpan.get(qn('w:val')))

        cell_borders = {}
        if tcPr is not None:
            tblBorders = tcPr.find(qn('w:tcBorders'))
            if tblBorders is not None:
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = tblBorders.find(qn(f'w:{border_name}'))
                    if border is not None:
                        cell_borders[border_name] = {
                            'val': border.get(qn('w:val')),
                            'color': border.get(qn('w:color')),
                            'sz': border.get(qn('w:sz')),
                        }

        shading = {}
        if tcPr is not None:
            shd = tcPr.find(qn('w:shd'))
            if shd is not None:
                shading = {
                    'fill': shd.get(qn('w:fill')),
                    'val': shd.get(qn('w:val')),
                }

        return {
            'rowspan': rowspan,
            'colspan': colspan,
            'borders': cell_borders,
            'shading': shading,
            'table_cols': len(table.columns),
            'table_rows': len(table.rows),
        }

    def _get_suggestions(self, text: str) -> List[str]:
        suggestions = []
        for region_type, keywords in REGION_TYPE_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text:
                    suggestions.append(region_type)
                    break
        if not suggestions:
            suggestions = ['教学目标', '教学重难点', '教学过程', '其他']
        return suggestions


class DocxFiller:
    def __init__(self, docx_bytes: bytes):
        self.docx_bytes = docx_bytes
        self.bio = BytesIO(docx_bytes)
        self.doc = Document(self.bio)

    def fill(self, content_mapping: Dict[str, str], structure_map: Dict[str, Dict]) -> BytesIO:
        for region_id, content in content_mapping.items():
            if region_id not in structure_map:
                continue

            region_info = structure_map[region_id]
            if region_info['type'] == 'paragraph':
                self._fill_paragraph(region_info, content)
            elif region_info['type'] == 'table_cell':
                self._fill_table_cell(region_info, content)
            elif region_info['type'] == 'textbox':
                self._fill_textbox(region_info, content)

        output_bio = BytesIO()
        self.doc.save(output_bio)
        output_bio.seek(0)
        # 修复 M5：填充并保存完成后，显式关闭输入 BytesIO，释放资源
        self.bio.close()
        return output_bio

    def _fill_paragraph(self, region_info: Dict, content: str):
        para_idx = region_info['index']
        if para_idx >= len(self.doc.paragraphs):
            return

        para = self.doc.paragraphs[para_idx]
        self._replace_paragraph_content(para, content)

    def _fill_table_cell(self, region_info: Dict, content: str):
        table_idx = region_info['table_index']
        row_idx = region_info['row_index']
        col_idx = region_info['col_index']

        if table_idx >= len(self.doc.tables):
            return

        table = self.doc.tables[table_idx]
        if row_idx >= len(table.rows):
            return

        row = table.rows[row_idx]
        if col_idx >= len(row.cells):
            return

        cell = row.cells[col_idx]

        # 在写入前记录原始容量
        original_capacity = self._estimate_cell_capacity(cell)

        estimated_size = len(content)
        if estimated_size > original_capacity * 1.5:
            # 溢出：在换行符边界截断，避免切断 Markdown 标记
            split_pos = self._find_safe_split_pos(content, original_capacity)
            first_half = content[:split_pos].rstrip()
            second_half = content[split_pos:].lstrip()
            self._replace_cell_content(cell, first_half)
            self._handle_cell_overflow(table, row_idx, col_idx, second_half, region_info, write_content=False)
        else:
            self._replace_cell_content(cell, content)

    def _find_safe_split_pos(self, content: str, target_pos: int) -> int:
        """在 target_pos 附近找到安全的截断点（换行符处），避免切断 Markdown 标记"""
        # 优先在换行符处截断
        newline_pos = content.rfind('\n', 0, target_pos + 20)
        if newline_pos > target_pos * 0.3:
            return newline_pos + 1

        # 其次在 ** 边界截断（确保不在加粗标记中间）
        bold_pos = content.rfind('**', 0, target_pos + 10)
        if bold_pos > target_pos * 0.3:
            # 确保 ** 是成对的
            count_before = content[:bold_pos].count('**')
            if count_before % 2 == 0:
                return bold_pos

        # 兜底：直接按 target_pos 截断
        return target_pos

    def _fill_textbox(self, region_info: Dict, content: str):
        """填充文本框内容：通过元素引用直接定位并替换文本"""
        if not content:
            return
        elem = region_info.get('_element_ref')
        if elem is None:
            return

        t_elements = list(elem.iter(qn('w:t')))
        if not t_elements:
            return

        # 修复 H2：先做 XML 转义，避免 < > & 破坏 XML 结构；按行拆分处理多行内容
        safe_content = xml_escape(str(content))
        lines = safe_content.split('\n')

        # 保留第一个 w:t 所在 run 的字体属性（rPr 不动），仅修改文本内容
        first_t = t_elements[0]
        first_t.text = lines[0] if lines else ''
        # 清空其余 w:t，避免残留旧文本
        for t in t_elements[1:]:
            t.text = ''

        # 多行内容：在第一个 run 内追加 <w:br/> 换行 + 后续行文本（复用该 run 字体属性）
        if len(lines) > 1:
            first_r = first_t.getparent()
            for line in lines[1:]:
                br = OxmlElement('w:br')
                first_r.append(br)
                new_t = OxmlElement('w:t')
                new_t.set(qn('xml:space'), 'preserve')
                new_t.text = line
                first_r.append(new_t)

    def _replace_paragraph_content(self, para, content: str):
        for run in list(para.runs):
            run._element.getparent().remove(run._element)
        lines = content.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                # 修复 H1：python-docx 中 \n 不会换行，需用 add_break() 插入换行符
                break_run = para.add_run()
                break_run.add_break()
            self._add_formatted_text(para, line, para.style.font)

    def _replace_cell_content(self, cell, content: str):
        original_font = None
        if cell.paragraphs and cell.paragraphs[0].runs and cell.paragraphs[0].runs[0]:
            original_font = cell.paragraphs[0].runs[0].font

        for para in cell.paragraphs:
            for run in list(para.runs):
                run._element.getparent().remove(run._element)

        if cell.paragraphs:
            first_para = cell.paragraphs[0]
        else:
            first_para = cell.add_paragraph()

        lines = content.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                # 修复 H1：python-docx 中 \n 不会换行，需用 add_break() 插入换行符
                break_run = first_para.add_run()
                break_run.add_break()
            self._add_formatted_text(first_para, line, original_font)

    def _add_formatted_text(self, para, line: str, source_font):
        """将一行文本按 Markdown 格式（**加粗**、*斜体*）拆分并添加到段落"""
        # 处理列表标记：- 或 1. 开头的行，添加缩进
        list_indent = ''
        stripped = line.lstrip()
        if stripped.startswith('- ') or stripped.startswith('• '):
            list_indent = '    '
            line = list_indent + '• ' + stripped[2:]
        elif len(stripped) >= 2 and stripped[0].isdigit() and stripped[1] == '.':
            list_indent = '    '
            line = list_indent + stripped

        # 按 ** 拆分处理加粗
        parts = line.split('**')
        if len(parts) == 1:
            # 没有加粗标记，直接处理斜体
            self._add_italic_text(para, line, source_font, bold=False)
            return

        for i, part in enumerate(parts):
            if not part:
                continue
            # 偶数索引是普通文本，奇数索引是加粗文本
            is_bold = (i % 2 == 1)
            self._add_italic_text(para, part, source_font, bold=is_bold)

    def _add_italic_text(self, para, text: str, source_font, bold: bool):
        """处理 *斜体* 标记"""
        if '*' not in text:
            run = para.add_run(text)
            self._apply_run_style(run, source_font)
            if bold:
                run.font.bold = True
            return

        parts = text.split('*')
        for i, part in enumerate(parts):
            if not part:
                continue
            run = para.add_run(part)
            self._apply_run_style(run, source_font)
            if bold:
                run.font.bold = True
            if i % 2 == 1:
                run.font.italic = True

    def _apply_run_style(self, run, source_font):
        if source_font is None:
            return

        run.font.name = source_font.name
        run.font.size = source_font.size
        run.font.bold = source_font.bold
        run.font.italic = source_font.italic
        if source_font.color:
            try:
                if source_font.color.rgb:
                    run.font.color.rgb = source_font.color.rgb
            except (ValueError, TypeError):
                pass

    def _estimate_cell_capacity(self, cell) -> int:
        text = cell.text
        # 修复 M1：对空单元格给定最小容量 50，避免容量估算为 0 导致逻辑失效
        base = len(text) * 3 if text else 0
        return max(base, 50)

    def _handle_cell_overflow(self, table: Table, row_idx: int, col_idx: int, content: str, region_info: Dict, write_content: bool = True, max_depth: int = 3):
        # 修复 M2：增加深度限制，防止递归处理二次溢出时无限循环
        if max_depth <= 0:
            return
        cell_info = region_info['cell_info']
        rowspan = cell_info.get('rowspan', 1)
        colspan = cell_info.get('colspan', 1)
        cell_borders = cell_info.get('borders', {})
        cell_shading = cell_info.get('shading', {})

        new_row = self._insert_row_below(table, row_idx)
        self._copy_row_styles(new_row, table.rows[row_idx])

        if colspan > 1:
            self._merge_cells_horizontally(new_row, col_idx, colspan)

        target_cell = new_row.cells[col_idx]
        self._apply_cell_styles(target_cell, cell_borders, cell_shading)

        if write_content:
            # 兼容旧调用方式：截断并写入两半
            mid = self._find_safe_split_pos(content, len(content) // 2)
            first_half = content[:mid].rstrip()
            second_half = content[mid:].lstrip()
            original_cell = table.rows[row_idx].cells[col_idx]
            self._replace_cell_content(original_cell, first_half)
            self._replace_cell_content(target_cell, second_half)
            # 修复 M2：递归处理二次溢出，防止 second_half 仍超出新单元格容量
            # 写入前 target_cell 为 deepcopy 的模板内容，据此估算容量判断是否需继续拆分
            new_cell_capacity = self._estimate_cell_capacity(target_cell)
            if max_depth > 1 and len(second_half) > new_cell_capacity:
                self._handle_cell_overflow(
                    table, row_idx + 1, col_idx, second_half,
                    region_info, write_content=True, max_depth=max_depth - 1
                )
        else:
            # 新模式：原单元格已写入，只需写入新行
            # 修复 M2：写入前先估算新单元格容量，用于判断 content（second_half）是否二次溢出
            new_cell_capacity = self._estimate_cell_capacity(target_cell)
            if max_depth > 1 and len(content) > new_cell_capacity:
                # content 仍溢出，递归拆分到新行（write_content=True 会把首部写回 target_cell）
                self._handle_cell_overflow(
                    table, row_idx + 1, col_idx, content,
                    region_info, write_content=True, max_depth=max_depth - 1
                )
            else:
                self._replace_cell_content(target_cell, content)

    def _insert_row_below(self, table: Table, row_idx: int) -> Row:
        tbl = table._tbl
        row_to_copy = table.rows[row_idx]._tr

        new_tr = deepcopy(row_to_copy)

        for tc in new_tr.findall(qn('w:tc')):
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                vMerge = tcPr.find(qn('w:vMerge'))
                if vMerge is not None:
                    tcPr.remove(vMerge)

        rows = tbl.findall(qn('w:tr'))
        if row_idx + 1 < len(rows):
            tbl.insert(row_idx + 1, new_tr)
        else:
            tbl.append(new_tr)

        # 修复 L1：不重置 python-docx 私有属性 _rows，依赖其对 w:tr 的惰性查询
        # （原代码: table._rows = [Row(r, table) for r in tbl.findall(qn('w:tr'))]）
        return table.rows[row_idx + 1]

    def _copy_row_styles(self, target_row: Row, source_row: Row):
        for target_cell, source_cell in zip(target_row.cells, source_row.cells):
            self._copy_cell_style(target_cell, source_cell)

    def _copy_cell_style(self, target_cell: Cell, source_cell: Cell):
        source_tc = source_cell._tc
        target_tc = target_cell._tc

        source_tcPr = source_tc.find(qn('w:tcPr'))
        if source_tcPr is not None:
            target_tcPr = target_tc.find(qn('w:tcPr'))
            if target_tcPr is not None:
                target_tc.remove(target_tcPr)
            new_tcPr = deepcopy(source_tcPr)
            target_tc.insert(0, new_tcPr)

        for source_para, target_para in zip(source_cell.paragraphs, target_cell.paragraphs):
            if source_para.style and target_para.style:
                target_para.style = source_para.style

    def _merge_cells_horizontally(self, row: Row, start_col: int, span: int):
        cells = row.cells
        if start_col + span > len(cells):
            span = len(cells) - start_col

        if span <= 1:
            return

        first_cell = cells[start_col]
        tcPr = first_cell._tc.get_or_add_tcPr()
        # 修复 M3：先移除已存在的 gridSpan，避免重复追加导致 XML 结构异常
        existing = tcPr.find(qn('w:gridSpan'))
        if existing is not None:
            tcPr.remove(existing)
        gridSpan = OxmlElement('w:gridSpan')
        gridSpan.set(qn('w:val'), str(span))
        tcPr.append(gridSpan)

        tr = first_cell._tc.getparent()
        tcs = tr.findall(qn('w:tc'))
        for i in range(start_col + 1, start_col + span):
            if i < len(tcs):
                tr.remove(tcs[i])

    def _apply_cell_styles(self, cell: Cell, borders: Dict, shading: Dict):
        tcPr = cell._tc.get_or_add_tcPr()

        if borders:
            tblBorders = OxmlElement('w:tcBorders')
            for border_name, border_props in borders.items():
                border = OxmlElement(f'w:{border_name}')
                if 'val' in border_props:
                    border.set(qn('w:val'), border_props['val'])
                if 'color' in border_props:
                    border.set(qn('w:color'), border_props['color'])
                if 'sz' in border_props:
                    border.set(qn('w:sz'), border_props['sz'])
                tblBorders.append(border)
            tcPr.append(tblBorders)

        if shading:
            shd = OxmlElement('w:shd')
            if 'fill' in shading:
                shd.set(qn('w:fill'), shading['fill'])
            if 'val' in shading:
                shd.set(qn('w:val'), shading['val'])
            tcPr.append(shd)


def parse_docx_structure(docx_bytes: bytes) -> Tuple[Dict, List[Dict]]:
    parser = DocxStructureParser(docx_bytes)
    return parser.parse()


def fill_lesson_plan_template(docx_bytes: bytes, content_mapping: Dict[str, str], structure_map: Dict[str, Dict]) -> BytesIO:
    filler = DocxFiller(docx_bytes)
    return filler.fill(content_mapping, structure_map)