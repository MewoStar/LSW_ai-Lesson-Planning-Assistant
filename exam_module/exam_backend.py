import os
import json
import re
import time
import uuid
import ipaddress
import traceback
import threading
from io import BytesIO
from html import escape
from urllib.parse import quote as _url_quote
from flask import request, jsonify, redirect, url_for, send_file
from openai import OpenAI


EXAM_MODULE_ID = "exam_paper"

_LOG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ai_response_logs")
os.makedirs(_LOG_DIR, exist_ok=True)

_EXAM_RATE_LIMIT = {}
_EXAM_RATE_LOCK = threading.Lock()

def _is_private_ip(ip):
    # 修复 M1：原 startswith 判断会把 172.x 全部当作私有，且会漏判 192.168、169.254 等。
    # 改用 ipaddress 模块，准确识别所有私有/环回/链路本地地址。
    try:
        return ipaddress.ip_address(ip).is_private
    except ValueError:
        return False


def _get_exam_client_ip():
    if request.headers.get('X-Forwarded-For'):
        forwarded = request.headers.get('X-Forwarded-For')
        parts = forwarded.split(',')
        for part in parts:
            ip = part.strip()
            # 修复 M1：用 ipaddress 判断私有 IP，替代原错误的 startswith 链
            if ip and not _is_private_ip(ip):
                return ip
        return parts[0].strip() if parts else 'unknown'
    return request.remote_addr or 'unknown'

def _check_exam_rate_limit(ip: str) -> bool:
    now = time.time()
    with _EXAM_RATE_LOCK:
        expired_keys = [k for k, v in _EXAM_RATE_LIMIT.items() if all(now - t > 60 for t in v)]
        for k in expired_keys:
            del _EXAM_RATE_LIMIT[k]
        if ip not in _EXAM_RATE_LIMIT:
            _EXAM_RATE_LIMIT[ip] = []
        attempts = _EXAM_RATE_LIMIT[ip]
        attempts = [t for t in attempts if now - t < 60]
        _EXAM_RATE_LIMIT[ip] = attempts
        if len(attempts) >= 5:
            return False
        _EXAM_RATE_LIMIT[ip].append(now)
    return True

def _cleanup_expired_files(output_dir, max_age_hours=24):
    try:
        if not os.path.exists(output_dir):
            return
        now = time.time()
        cutoff = now - max_age_hours * 3600
        for filename in os.listdir(output_dir):
            filepath = os.path.join(output_dir, filename)
            if os.path.isfile(filepath):
                try:
                    mtime = os.path.getmtime(filepath)
                    if mtime < cutoff and (filename.endswith('.docx') or filename.endswith('.pdf')):
                        os.remove(filepath)
                except Exception as cleanup_e:
                    # 修复 M6：原 except 完全静默，改为至少打印异常便于排查
                    print(f"清理临时文件失败: {cleanup_e}")
    except Exception as cleanup_e:
        # 修复 M6：外层异常也打印，避免完全静默
        print(f"清理临时文件失败: {cleanup_e}")


def register_exam_routes(
    app,
    get_current_user_func,
    database_mod,
    api_key,
    base_url,
    model_name,
    get_user_output_dir_func,
):
    """注册试卷生成模块所有路由。

    参数与项目现有的 register 模式保持一致，依赖从主应用注入。
    """

    get_current_user = get_current_user_func
    database = database_mod
    get_user_output_dir = get_user_output_dir_func

    @app.route("/exam")
    def exam_redirect():
        return redirect(url_for("module_detail", module_id=EXAM_MODULE_ID))

    @app.route("/api/exam/generate", methods=["POST"])
    def api_exam_generate():
        user = get_current_user()
        # 修复 H1：未登录用户原匿名放行，存在滥用风险，改为要求登录返回 401
        if not user:
            return jsonify({"success": False, "error": "请先登录"}), 401

        try:
            ip = _get_exam_client_ip()

            if not _check_exam_rate_limit(ip):
                return jsonify({"success": False, "error": "请求过于频繁，请1分钟后再试"}), 429

            data = request.get_json() or {}

            def _safe_int(v, default=0):
                try:
                    if v is None:
                        return default
                    return int(v)
                except (TypeError, ValueError):
                    return default

            # 修复 H2：原 exam_title 直接 [:100] 截断，导致后续 len>100 校验永远为 False，校验失效。
            # 改为先 strip 不截断，让长度校验生效返回错误，由前端提示用户修改。
            exam_title = (data.get("exam_title") or "").strip()
            topic = (data.get("topic") or "").strip()[:500]
            subject = (data.get("subject") or "").strip()[:50]
            grade = (data.get("grade") or "").strip()[:50]
            textbook = (data.get("textbook") or "").strip()[:100]
            difficulty = (data.get("difficulty") or "中等").strip()[:20]
            total_score = max(1, min(_safe_int(data.get("total_score", 100), 100), 1000))
            duration = max(1, min(_safe_int(data.get("duration", 90), 90), 300))
            question_types = data.get("question_types", {}) or {}

            valid_qtypes = {"choice", "fill", "judge", "answer"}
            if isinstance(question_types, dict):
                question_types = {k: v for k, v in question_types.items() if k in valid_qtypes}

            def _get_qtype_count(val):
                if isinstance(val, dict):
                    return _safe_int(val.get('count', 0), 0)
                return _safe_int(val, 0)

            def _get_qtype_score(val):
                if isinstance(val, dict):
                    return _safe_int(val.get('score', 0), 0)
                return 0

            if not topic:
                return jsonify({"success": False, "error": "请输入知识点"}), 400

            if exam_title and (len(exam_title) > 100 or re.search(r'[<>:"/\\|?*]', exam_title)):
                return jsonify({"success": False, "error": "试卷标题不能超过100字符，且不能包含特殊字符"}), 400

            total_questions = 0
            if isinstance(question_types, dict) and question_types:
                total_questions = sum(_get_qtype_count(item) for item in question_types.values())
            if total_questions > 50:
                return jsonify({"success": False, "error": f"题目总数({total_questions}道)超过上限(50道)，请减少题目数量"}), 400
            if total_questions == 0:
                return jsonify({"success": False, "error": "请至少选择一种题型并设置数量"}), 400

            if not api_key:
                return jsonify({"success": False, "error": "未配置 API Key"}), 500

            calculated_total = 0
            for qt, val in question_types.items():
                cnt = _get_qtype_count(val)
                sc = _get_qtype_score(val)
                if cnt > 0 and sc > 0:
                    calculated_total += cnt * sc

            # 修复 M3：原代码静默用 calculated_total 覆盖 total_score，用户无感知。
            # 改为保留覆盖逻辑，但同时生成警告字段，最终返回给前端提示用户。
            warning = None
            if calculated_total > 0 and calculated_total != total_score:
                print(f"[exam] 总分不一致：用户指定{total_score}分，题型配置计算为{calculated_total}分，使用配置值")
                warning = f"题型分值合计为 {calculated_total} 分，与指定总分 {total_score} 分不一致，已按实际分值生成"
                total_score = calculated_total

            qtype_names = {
                "choice": "选择题",
                "fill": "填空题",
                "judge": "判断题",
                "answer": "解答题",
            }
            qtype_config = []
            for qt, val in question_types.items():
                cnt = _get_qtype_count(val)
                if cnt > 0:
                    sc = _get_qtype_score(val)
                    if sc > 0:
                        qtype_config.append(f"{qtype_names.get(qt, qt)}：{cnt}题（每题{sc}分）")
                    else:
                        qtype_config.append(f"{qtype_names.get(qt, qt)}：{cnt}题")

            qtype_details = []
            for qt, val in question_types.items():
                cnt = _get_qtype_count(val)
                if cnt > 0:
                    qtype_details.append({
                        "type": qtype_names.get(qt, qt),
                        "count": cnt,
                        "score": _get_qtype_score(val),
                        "key": qt
                    })

            prompt = f"""请生成一份{subject}{grade}试卷，知识点：{topic}，难度：{difficulty}，总分{total_score}分，时长{duration}分钟。

题型配置：{', '.join(qtype_config)}

输出格式：直接输出JSON，不要任何其他文字。
{{
    "title": "试卷名称",
    "subject": "{subject}",
    "grade": "{grade}",
    "total_score": {total_score},
    "duration": {duration},
    "difficulty": "{difficulty}",
    "sections": [
        {{
            "type": "题型名称",
            "count": 数量,
            "questions": [
                {{
                    "id": 题号,
                    "content": "题目内容",
                    "options": ["A选项","B选项","C选项","D选项"],
                    "answer": "答案",
                    "score": 分值,
                    "knowledge_point": "知识点",
                    "difficulty_label": "简单/中等/困难",
                    "analysis": "解析"
                }}
            ]
        }}
    ],
    "answer_sheet": "答案列表"
}}

规则：选择题options为4个选项，answer为A/B/C/D；填空题和解答题options为[]；判断题answer为"正确"/"错误"；题号连续；总分等于各题分值之和。"""

            params_summary = f"subject={subject}, grade={grade}, topic={topic[:30]}, total_score={total_score}, questions={total_questions}"
            print(f"[exam] Starting generation: {params_summary}")

            # TODO: 应复用客户端连接池，避免每次请求新建 OpenAI 客户端造成连接开销。
            # 后续可在 register_exam_routes 闭包外创建一次 client 并捕获复用。
            client = OpenAI(api_key=api_key, base_url=base_url, timeout=180.0)
            response = client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": "你是专业的教师考试命题专家，只输出 JSON 格式的试卷数据。"},
                    {"role": "user", "content": prompt},
                ],
                temperature=0.7,
                max_tokens=8192,
            )

            if not response.choices or not response.choices[0].message:
                raise ValueError("AI 返回为空")

            reply = response.choices[0].message.content

            print(f"[exam] AI response received: choices={len(response.choices)}, content_length={len(reply) if reply else 0}")

            _save_ai_response_log(reply, params_summary)

            exam_json = _parse_exam_json(reply)

            if exam_json.get("title") == "试卷生成失败":
                return jsonify({"success": False, "error": "AI 返回的数据格式异常，请重试或检查知识点描述"}), 500

            result = {
                "success": True,
                "message": "试卷生成成功",
                "topic": topic,
                "exam_data": exam_json,
            }
            # 修复 M3：如有总分不一致警告，加入响应字段由前端提示
            if warning:
                result["warning"] = warning

            return jsonify(result)

        except ValueError as ve:
            print(f"[exam] ValueError: {ve}")
            return jsonify({"success": False, "error": f"生成失败：{str(ve)}"}), 500
        except Exception as e:
            traceback.print_exc()
            error_str = str(e).lower()
            if 'timeout' in error_str or 'timed out' in error_str:
                return jsonify({"success": False, "error": "请求超时，请稍后重试"}), 504
            elif 'api key' in error_str or 'invalid key' in error_str or 'authentication' in error_str:
                return jsonify({"success": False, "error": "API Key 无效或未配置，请联系管理员"}), 500
            elif 'rate limit' in error_str or 'quota' in error_str:
                return jsonify({"success": False, "error": "API 请求配额不足，请稍后重试"}), 429
            else:
                return jsonify({"success": False, "error": f"生成失败：{str(e)[:100]}"}), 500

    @app.route("/api/exam/export/docx", methods=["POST"])
    def api_exam_export_docx():
        user = get_current_user()
        # 修复 H1：未登录用户原匿名放行，改为要求登录返回 401
        if not user:
            return jsonify({"success": False, "error": "请先登录"}), 401

        try:
            data = request.get_json() or {}
            exam_data = data.get("exam_data")
            export_type = data.get("type", "exam")

            if not exam_data:
                return jsonify({"success": False, "error": "试卷数据为空"}), 400

            user_output_dir = get_user_output_dir(user['id'])
            _cleanup_expired_files(user_output_dir)

            if export_type == "answer":
                buffer = _exam_to_word_answer(exam_data)
                filename_suffix = "_答案解析"
            else:
                buffer = _exam_to_word(exam_data)
                filename_suffix = ""

            timestamp = time.strftime("%Y%m%d_%H%M%S")
            safe_title = re.sub(r'[<>:"/\\|?*]', '', exam_data.get("title", "试卷"))[:20] or "试卷"
            filename = f"{safe_title}{filename_suffix}_{timestamp}.docx"
            filepath = os.path.join(user_output_dir, filename)

            with open(filepath, "wb") as f:
                f.write(buffer.getvalue())

            # 修复 M2：原代码 encode('ascii', errors='ignore') 会丢失中文字符，导致下载文件名变成纯数字时间戳。
            # 改为保留中文文件名，Flask 的 send_file 会自动按 RFC 5987 编码 Content-Disposition。
            # 同时提供 ASCII 兜底名（仅用于 X-Filename 头，便于不支持中文的环境降级）。
            try:
                _ascii_fn = filename.encode('ascii', errors='ignore').decode('ascii') or 'exam.docx'
            except Exception:
                _ascii_fn = 'exam.docx'

            resp = send_file(filepath, as_attachment=True, download_name=filename,
                             mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            resp.headers['X-Filename-Encoded'] = _url_quote(filename)
            resp.headers['X-Filename'] = _ascii_fn
            resp.headers['Access-Control-Expose-Headers'] = 'X-Filename, X-Filename-Encoded, Content-Disposition'
            return resp

        except Exception as e:
            traceback.print_exc()
            return jsonify({"success": False, "error": "导出失败，请稍后重试"}), 500

    @app.route("/api/exam/export/pdf", methods=["POST"])
    def api_exam_export_pdf():
        user = get_current_user()
        # 修复 H1：未登录用户原匿名放行，改为要求登录返回 401
        if not user:
            return jsonify({"success": False, "error": "请先登录"}), 401

        # 修复 M4：默认值预先初始化，防止 try 内抛异常后降级逻辑引用 export_type 时 NameError
        export_type = "exam"

        try:
            data = request.get_json() or {}
            exam_data = data.get("exam_data")
            export_type = data.get("type", "exam")

            if not exam_data:
                return jsonify({"success": False, "error": "试卷数据为空"}), 400

            user_output_dir = get_user_output_dir(user['id'])
            _cleanup_expired_files(user_output_dir)

            if export_type == "answer":
                html_content = _exam_to_html_answer(exam_data)
                filename_suffix = "_答案解析"
            else:
                html_content = _exam_to_html(exam_data)
                filename_suffix = ""

            buffer = _html_to_pdf(html_content)

            timestamp = time.strftime("%Y%m%d_%H%M%S")
            safe_title = re.sub(r'[<>:"/\\|?*]', '', exam_data.get("title", "试卷"))[:20] or "试卷"
            filename = f"{safe_title}{filename_suffix}_{timestamp}.pdf"
            filepath = os.path.join(user_output_dir, filename)

            with open(filepath, "wb") as f:
                f.write(buffer.getvalue())

            # 修复 M2：保留中文文件名，send_file 会自动按 RFC 5987 编码
            try:
                _ascii_fn = filename.encode('ascii', errors='ignore').decode('ascii') or 'exam.pdf'
            except Exception:
                _ascii_fn = 'exam.pdf'

            resp = send_file(filepath, as_attachment=True, download_name=filename,
                             mimetype='application/pdf')
            resp.headers['X-Filename-Encoded'] = _url_quote(filename)
            resp.headers['X-Filename'] = _ascii_fn
            resp.headers['Access-Control-Expose-Headers'] = 'X-Filename, X-Filename-Encoded, Content-Disposition'
            return resp

        except Exception as e:
            traceback.print_exc()
            pdf_error = str(e)
            if 'weasyprint' in pdf_error.lower() or 'libgobject' in pdf_error.lower() or 'GTK' in pdf_error:
                try:
                    user_output_dir = get_user_output_dir(user['id'])
                    if export_type == "answer":
                        buffer = _exam_to_word_answer(exam_data)
                        filename_suffix = "_答案解析"
                        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        ext = '.docx'
                    else:
                        buffer = _exam_to_word(exam_data)
                        filename_suffix = ""
                        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        ext = '.docx'

                    timestamp = time.strftime("%Y%m%d_%H%M%S")
                    safe_title = re.sub(r'[<>:"/\\|?*]', '', exam_data.get("title", "试卷"))[:20] or "试卷"
                    filename = f"{safe_title}{filename_suffix}_{timestamp}{ext}"
                    filepath = os.path.join(user_output_dir, filename)

                    with open(filepath, "wb") as f:
                        f.write(buffer.getvalue())

                    # 修复 M2：保留中文文件名，send_file 会自动按 RFC 5987 编码
                    try:
                        _ascii_fn = filename.encode('ascii', errors='ignore').decode('ascii') or f'exam{ext}'
                    except Exception:
                        _ascii_fn = f'exam{ext}'

                    resp = send_file(filepath, as_attachment=True, download_name=filename,
                                     mimetype=mime_type)
                    resp.headers['X-Filename-Encoded'] = _url_quote(filename)
                    resp.headers['X-Filename'] = _ascii_fn
                    resp.headers['X-Filename-Downgraded'] = 'pdf_to_docx'
                    resp.headers['Access-Control-Expose-Headers'] = 'X-Filename, X-Filename-Encoded, X-Filename-Downgraded, Content-Disposition'
                    return resp
                except Exception as downgrade_e:
                    # 修复 M6：原 except 完全静默，改为打印异常便于排查降级失败原因
                    print(f"PDF 降级 docx 导出失败: {downgrade_e}")
                    return jsonify({"success": False, "error": "导出失败，请稍后重试"}), 500
            return jsonify({"success": False, "error": "导出失败，请稍后重试"}), 500

    return app


def _save_ai_response_log(reply_text, params_summary=""):
    try:
        ts = time.strftime("%Y%m%d_%H%M%S")
        # 修复 M5：原文件名仅精确到秒，并发或快速重试时易发生同名覆盖。追加 uuid 前 6 位避免冲突。
        filename = f"ai_response_{ts}_{uuid.uuid4().hex[:6]}.txt"
        filepath = os.path.join(_LOG_DIR, filename)
        content = f"=== AI Response Log [{ts}] ===\n"
        content += f"Params: {params_summary}\n"
        content += f"Reply length: {len(reply_text) if reply_text else 0}\n"
        content += "=== Reply Start ===\n"
        content += (reply_text or "") + "\n"
        content += "=== Reply End ===\n"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        latest_path = os.path.join(_LOG_DIR, "latest_reply.txt")
        with open(latest_path, "w", encoding="utf-8") as f:
            f.write(content)
        print(f"[exam] AI response logged to: {filepath}")
    except Exception as e:
        print(f"[exam] Failed to save log: {e}")


def _extract_json(reply_text):
    # 修复 H3：原 re.search(r"\{[\s\S]*\}") 贪婪匹配会匹配到最后一个 }，
    # 当回复中存在多余 }（如示例 JSON 块后的说明文字）时会导致 json.loads 失败。
    # 改为优先用 JSONDecoder().raw_decode 从首个 { 起解析，仅解析有效 JSON 前缀，
    # 失败时再兜底尝试贪婪匹配。
    if not reply_text:
        return None
    text = reply_text.strip()
    # 去除可能的 markdown 代码块标记（```json ... ``` 或 ``` ... ```）
    if text.startswith('```'):
        lines = text.split('\n')
        lines = [l for l in lines if not l.strip().startswith('```')]
        text = '\n'.join(lines)
    # 找到第一个 {
    start = text.find('{')
    if start == -1:
        return None
    try:
        obj, _ = json.JSONDecoder().raw_decode(text[start:])
        return obj
    except Exception:
        # 兜底：尝试贪婪匹配最后一个 }
        m = re.search(r"\{[\s\S]*\}", text[start:])
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                return None
    return None


def _parse_exam_json(reply_text):
    if not reply_text:
        print("[exam] AI 返回内容为空")
        return {
            "title": "试卷生成失败",
            "subject": "",
            "grade": "",
            "total_score": 100,
            "duration": 90,
            "difficulty": "中等",
            "sections": [],
            "answer_sheet": "AI 返回内容为空，请重试",
        }

    try:
        match = re.search(r"```json\s*([\s\S]*?)\s*```", reply_text)
        if match:
            json_str = match.group(1).strip()
            print(f"[exam] Found code block JSON, length: {len(json_str)}")
            try:
                return json.loads(json_str)
            except json.JSONDecodeError:
                pass

        # 修复 H3：替换原贪婪正则匹配，改用更稳健的 _extract_json
        obj = _extract_json(reply_text)
        if obj is not None:
            print(f"[exam] Extracted JSON via _extract_json, type={type(obj).__name__}")
            return obj

        try:
            return json.loads(reply_text)
        except json.JSONDecodeError:
            pass

    except json.JSONDecodeError as e:
        print(f"[exam] JSON 解析失败: {e}")
        print(f"[exam] AI 返回内容前500字符: {reply_text[:500] if reply_text else '空'}")
    except Exception as e:
        print(f"[exam] 解析异常: {e}")

    return {
        "title": "试卷生成失败",
        "subject": "",
        "grade": "",
        "total_score": 100,
        "duration": 90,
        "difficulty": "中等",
        "sections": [],
        "answer_sheet": "AI 返回的数据格式异常，请重试或检查知识点描述",
    }


def _exam_to_word(exam_data):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise Exception("需要安装 python-docx 库，请在项目目录执行：pip install python-docx")

    doc = Document()
    for section in doc.sections:
        from docx.shared import Cm
        section.page_width = Cm(29.7)
        section.page_height = Cm(42.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(exam_data.get("title", "试卷"))
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_line = f"学科：{exam_data.get('subject', '')} | 年级：{exam_data.get('grade', '')} | 总分：{exam_data.get('total_score', 100)}分 | 时长：{exam_data.get('duration', 90)}分钟 | 难度：{exam_data.get('difficulty', '中等')}"
    info_para = doc.add_paragraph(info_line)
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.runs[0].font.size = Pt(10)
    info_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    q_num = 1
    for section in exam_data.get("sections", []):
        sec_heading = doc.add_heading(f"{section.get('type', '')}（共{section.get('count', 0)}题）", level=1)
        sec_heading.runs[0].font.size = Pt(14)
        sec_heading.runs[0].font.color.rgb = RGBColor(31, 78, 121)

        for q in section.get("questions", []):
            q_para = doc.add_paragraph()
            q_para.add_run(f"{q_num}. ").bold = True
            q_para.add_run(q.get("content", ""))

            if q.get("options") and len(q.get("options")) > 0:
                for idx, opt in enumerate(q.get("options")):
                    opt_para = doc.add_paragraph()
                    opt_para.paragraph_format.left_indent = Pt(20)
                    opt_para.add_run(f"{chr(65 + idx)}. ").bold = True
                    opt_para.add_run(opt)

            meta_line = f"【知识点】{q.get('knowledge_point', '')} | 【难度】{q.get('difficulty_label', '')} | 【分值】{q.get('score', 0)}分"
            meta_para = doc.add_paragraph(meta_line)
            meta_para.paragraph_format.left_indent = Pt(10)
            meta_para.runs[0].font.size = Pt(9)
            meta_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)

            doc.add_paragraph()
            q_num += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _exam_to_word_answer(exam_data):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise Exception("需要安装 python-docx 库，请在项目目录执行：pip install python-docx")

    doc = Document()
    for section in doc.sections:
        from docx.shared import Cm
        section.page_width = Cm(29.7)
        section.page_height = Cm(42.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{exam_data.get('title', '试卷')} - 参考答案及解析")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    answer_sheet = exam_data.get("answer_sheet", "")
    if answer_sheet and "AI 返回的数据格式异常" not in answer_sheet:
        sheet_heading = doc.add_heading("答案速查表", level=2)
        sheet_heading.runs[0].font.size = Pt(14)
        doc.add_paragraph(answer_sheet)
        doc.add_paragraph()

    q_num = 1
    for section in exam_data.get("sections", []):
        sec_heading = doc.add_heading(f"{section.get('type', '')}", level=1)
        sec_heading.runs[0].font.size = Pt(14)

        for q in section.get("questions", []):
            q_para = doc.add_paragraph()
            q_para.add_run(f"{q_num}. ").bold = True
            q_para.add_run(q.get("content", ""))

            ans_para = doc.add_paragraph()
            ans_para.add_run("答案：").bold = True
            ans_run = ans_para.add_run(q.get("answer", ""))
            ans_run.font.color.rgb = RGBColor(0, 128, 0)
            ans_run.bold = True

            if q.get("knowledge_point"):
                kp_para = doc.add_paragraph()
                kp_para.add_run("知识点：").bold = True
                kp_para.add_run(q.get("knowledge_point"))

            if q.get("difficulty_label"):
                diff_para = doc.add_paragraph()
                diff_para.add_run("难度：").bold = True
                diff_para.add_run(q.get("difficulty_label"))

            if q.get("analysis"):
                ana_heading = doc.add_heading("解析：", level=3)
                ana_heading.runs[0].font.size = Pt(11)
                doc.add_paragraph(q.get("analysis"))

            doc.add_paragraph()
            q_num += 1

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _exam_to_html(exam_data):
    title = escape(exam_data.get('title', '试卷'))
    subject = escape(exam_data.get('subject', ''))
    grade = escape(exam_data.get('grade', ''))
    total_score = escape(str(exam_data.get('total_score', 100)))
    duration = escape(str(exam_data.get('duration', 90)))
    difficulty = escape(exam_data.get('difficulty', '中等'))

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
    @page {{ size: A3; margin: 2cm 2.5cm; }}
    body {{ font-family: 'Microsoft YaHei', sans-serif; font-size: 14px; line-height: 1.8; padding: 40px; }}
    h1 {{ text-align: center; font-size: 22px; color: #1f4e79; border-bottom: 2px solid #1f4e79; padding-bottom: 10px; }}
    .info {{ text-align: center; color: #666; font-size: 13px; margin-bottom: 20px; }}
    h2 {{ font-size: 15px; color: #1f4e79; margin-top: 24px; }}
    .question {{ margin: 16px 0; }}
    .q-num {{ font-weight: bold; }}
    .options {{ margin-left: 24px; }}
    .option {{ margin: 8px 0; }}
    .meta {{ font-size: 12px; color: #999; margin-top: 8px; }}
    .diff-easy {{ color: #065f46; }}
    .diff-medium {{ color: #92400e; }}
    .diff-hard {{ color: #991b1b; }}
</style>
</head>
<body>
<h1>{title}</h1>
<p class="info">学科：{subject} | 年级：{grade} | 总分：{total_score}分 | 时长：{duration}分钟 | 难度：{difficulty}</p>
"""
    q_num = 1
    for section in exam_data.get("sections", []):
        sec_type = escape(section.get('type', ''))
        html += f"<h2>{sec_type}（共{section.get('count', 0)}题）</h2>"
        for q in section.get("questions", []):
            q_content = escape(q.get("content", ""))
            html += f'<div class="question"><span class="q-num">{q_num}.</span>{q_content}'
            if q.get("options") and len(q.get("options")) > 0:
                html += '<div class="options">'
                for idx, opt in enumerate(q.get("options")):
                    opt_text = escape(opt)
                    html += f'<div class="option">{chr(65 + idx)}. {opt_text}</div>'
                html += '</div>'
            kp = escape(q.get("knowledge_point", ""))
            diff_label = escape(q.get("difficulty_label", ""))
            diff_class = 'diff-easy' if diff_label == '简单' else \
                         'diff-hard' if diff_label == '困难' else 'diff-medium'
            html += f'<div class="meta">知识点：{kp} | <span class="{diff_class}">难度：{diff_label}</span> | 分值：{q.get("score", 0)}分</div></div>'
            q_num += 1

    html += "</body></html>"
    return html


def _exam_to_html_answer(exam_data):
    title = escape(exam_data.get('title', '试卷'))

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{title} - 答案解析</title>
<style>
    @page {{ size: A3; margin: 2cm 2.5cm; }}
    body {{ font-family: 'Microsoft YaHei', sans-serif; font-size: 14px; line-height: 1.8; padding: 40px; }}
    h1 {{ text-align: center; font-size: 22px; color: #1f4e79; }}
    h2 {{ font-size: 15px; color: #1f4e79; margin-top: 20px; }}
    h3 {{ font-size: 13px; color: #008000; }}
    .answer {{ color: #008000; font-weight: bold; }}
    .analysis {{ color: #666; }}
    .diff-easy {{ color: #065f46; }}
    .diff-medium {{ color: #92400e; }}
    .diff-hard {{ color: #991b1b; }}
</style>
</head>
<body>
<h1>{title} - 参考答案及解析</h1>
"""
    answer_sheet = exam_data.get("answer_sheet", "")
    if answer_sheet and "AI 返回的数据格式异常" not in answer_sheet:
        html += '<h2>答案速查表</h2><p>' + escape(answer_sheet) + '</p>'

    q_num = 1
    for section in exam_data.get("sections", []):
        sec_type = escape(section.get('type', ''))
        html += f"<h2>{sec_type}</h2>"
        for q in section.get("questions", []):
            q_content = escape(q.get("content", ""))
            q_answer = escape(q.get("answer", ""))
            html += f"<p><strong>{q_num}.</strong> {q_content}</p>"
            html += f"<p><span class='answer'>答案：{q_answer}</span></p>"
            if q.get("knowledge_point"):
                kp = escape(q.get("knowledge_point"))
                html += f"<p><strong>知识点：</strong>{kp}</p>"
            diff_label = escape(q.get("difficulty_label", ""))
            diff_class = 'diff-easy' if diff_label == '简单' else \
                         'diff-hard' if diff_label == '困难' else 'diff-medium'
            if diff_label:
                html += f"<p><strong>难度：</strong><span class='{diff_class}'>{diff_label}</span></p>"
            if q.get("analysis"):
                analysis = escape(q.get("analysis"))
                html += f"<h3>解析：</h3><p class='analysis'>{analysis}</p>"
            html += "<hr>"
            q_num += 1

    html += "</body></html>"
    return html


def _html_to_pdf(html_content):
    try:
        from weasyprint import HTML
    except (ImportError, OSError):
        raise Exception("PDF导出需要 weasyprint 库及 GTK 运行时环境。请安装 GTK3 后重试，或使用 Word 导出功能。")

    try:
        buffer = BytesIO()
        HTML(string=html_content).write_pdf(buffer)
        buffer.seek(0)
        return buffer
    except OSError as e:
        if 'libgobject' in str(e) or 'library' in str(e).lower():
            raise Exception("PDF导出需要安装 GTK3 运行时环境（libgobject-2.0-0）。请安装 GTK3 或使用 Word 导出功能。")
        raise Exception(f"PDF生成失败: {str(e)}")
